from flask import Flask, request, send_file, render_template, redirect, url_for
from docx import Document
from io import BytesIO
from docx.shared import Inches
import os
import vertexai
from vertexai.generative_models import GenerativeModel


# Read application variables from the config fle
BOTNAME = "Language Translator"
SUBTITLE = "Your Language Translator"
# Initialize Flask application
app = Flask(__name__)

# Define project information
PROJECT_ID = "vertexai project id"  # @param {type:"string"}
#LOCATION = "us-east4"  # @param {type:"string"}

# Initialize Vertex AI
import vertexai

vertexai.init(project=PROJECT_ID)

def translate_text_gemini(text, source_language="en", target_language="ja"):
    """
    Translates text using Vertex AI's gemini-1.0-pro model.
    """
    try:
        model = GenerativeModel("gemini-1.5-pro-001")
        prompt = f"Translate the following text from {source_language} to {target_language}:\n{text}"
        response = model.generate_content(prompt)
        return response.text if response else text
    except Exception as e:
        print(f"Translation error: {e}")
        return text

def translate_document(file, source_language="en", target_language="ja"):
    """
    Reads and translates a Word document, preserving images.
    """
    doc = Document(file)
    translated_doc = Document()

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Translate non-empty paragraphs
            translated_text = translate_text_gemini(paragraph.text, source_language, target_language)
            translated_doc.add_paragraph(translated_text)
        else:
            translated_doc.add_paragraph("")  # Preserve empty paragraphs

    # Copy images from the original document to the translated one
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_stream = rel.target_part.blob
            new_paragraph = translated_doc.add_paragraph()
            new_paragraph.add_run().add_picture(BytesIO(image_stream), width=Inches(1.0))

    # Save the translated document to a binary stream
    output_stream = BytesIO()
    translated_doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

@app.route("/", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        if "file" not in request.files:
            return "No file uploaded!", 400
        file = request.files["file"]
        if file.filename == "":
            return "No selected file!", 400
        if file:
            translated_file = translate_document(file)
            return send_file(
                translated_file,
                as_attachment=True,
                download_name="translated_document.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    
    # Display the home page with the required variables set
    answer = "Need English to Any Language translations? Leverage Gemini's expertise! I provide English-to-Any document, audio and video translation services."
    config = {
        "title": BOTNAME,
        #"subtitle": SUBTITLE,
        "botname": BOTNAME,
        "message": answer,
    }

    return render_template("index.html", config=config)

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))
